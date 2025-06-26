import pandas as pd
import streamlit as st
import psycopg2
import ast
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
from io import BytesIO
import unicodedata
import requests
from PIL import Image, ImageDraw
import io
import re

st.set_page_config(page_title="Archetypy Krzysztofa Hetmana – panel administratora", layout="wide")

ARCHE_NAMES_ORDER = [
    "Niewinny", "Mędrzec", "Odkrywca", "Buntownik", "Czarodziej", "Bohater",
    "Kochanek", "Błazen", "Towarzysz", "Opiekun", "Władca", "Twórca"
]

archetypes = {
    "Władca":   [1, 2, 3, 4],
    "Bohater":  [5, 6, 7, 8],
    "Mędrzec":  [9, 10, 11, 12],
    "Opiekun":  [13, 14, 15, 16],
    "Kochanek": [17, 18, 19, 20],
    "Błazen":   [21, 22, 23, 24],
    "Twórca":   [25, 26, 27, 28],
    "Odkrywca": [29, 30, 31, 32],
    "Czarodziej": [33, 34, 35, 36],
    "Towarzysz": [37, 38, 39, 40],
    "Niewinny": [41, 42, 43, 44],
    "Buntownik": [45, 46, 47, 48],
}

archetype_features = {
    "Władca": "Potrzeba kontroli, organizacji, zarządzanie, wprowadzanie ładu.",
    "Bohater": "Odwaga, walka z przeciwnościami, mobilizacja do działania.",
    "Mędrzec": "Wiedza, analityczność, logiczne argumenty, racjonalność.",
    "Opiekun": "Empatia, dbanie o innych, ochrona, troska.",
    "Kochanek": "Relacje, emocje, bliskość, autentyczność uczuć.",
    "Błazen": "Poczucie humoru, dystans, lekkość, rozładowywanie napięć.",
    "Twórca": "Kreatywność, innowacja, wyrażanie siebie, estetyka.",
    "Odkrywca": "Niezależność, zmiany, nowe doświadczenia, ekspresja.",
    "Czarodziej": "Transformacja, inspiracja, zmiana świata, przekuwanie idei w czyn.",
    "Towarzysz": "Autentyczność, wspólnota, prostota, bycie częścią grupy.",
    "Niewinny": "Optymizm, ufność, unikanie konfliktów, pozytywne nastawienie.",
    "Buntownik": "Kwestionowanie norm, odwaga w burzeniu zasad, radykalna zmiana."
}
# W TYM MIEJSCU WSTAW SWÓJ SŁOWNIK archetype_extended = {...}

archetype_extended = {
    "Władca": {
        "name": "Władca",
        "tagline": "Autorytet. Kontrola. Doskonałość.",
        "description": (
            "Archetyp Władcy w polityce uosabia siłę przywództwa, stabilność i pewność działania. "
            "Jako kandydat na prezydenta Lublina Władca stawia na porządek, wyznaczanie standardów rozwoju i podejmowanie stanowczych decyzji dla dobra wspólnego. "
            "Jest symbolem autentycznego autorytetu, przewodzenia i skutecznego zarządzania miastem. "
            "Buduje zaufanie, komunikując skuteczność, odpowiedzialność i gwarantując bezpieczeństwo mieszkańcom."
        ),
        "storyline": (
            "Narracja kampanii oparta na Władcy podkreśla spójność działań, panowanie nad trudnymi sytuacjami i sprawność w zarządzaniu miastem. "
            "Władca nie podąża za modą – wyznacza nowe standardy w samorządzie. "
            "Akcentuje dokonania, referencje i doświadczenie. Buduje obraz lidera odpowiadającego za przyszłość i prestiż Lublina."
        ),
        "recommendations": [
            "Używaj kolorystyki kojarzącej się z autorytetem – czerń, złoto, ciemny granat, burgund.",
            "Projektuj symbole: sygnety, herby miasta Lublin, podkreślając prestiż i zarządzanie.",
            "Komunikuj się językiem odpowiedzialności i troski o przyszłość miasta.",
            "Przekazuj komunikaty stanowczo, jednoznacznie, jako gospodarz miasta.",
            "Pokazuj osiągnięcia, inwestycje, referencje mieszkańców.",
            "Zadbaj o trwałość i jakość działań – nie obniżaj standardów.",
            "Twórz aurę elitarności: zamknięte konsultacje, spotkania liderów opinii.",
            "Przyciągaj wyborców ceniących bezpieczeństwo, stabilizację i prestiż miasta.",
            "Unikaj luźnego, żartobliwego tonu – postaw na klasę i profesjonalizm."
        ],
        "core_traits": [
            "Autorytet", "Przywództwo", "Stabilność", "Prestiż", "Kontrola", "Inspiracja", "Mistrzostwo"
        ],
        "example_brands": [
            "Rolex", "Mercedes-Benz", "IBM", "British Airways", "Silny samorząd"
        ],
        "color_palette": [
            "#000000", "#FFD700", "#282C34", "#800020"
        ],
        "visual_elements": [
            "Korona", "Herb Lublina", "Sygnet", "Monogram", "Geometryczna, masywna typografia", "Symetria"
        ],
        "keyword_messaging": [
            "Lider Lublina", "Siła samorządu", "Stabilność", "Doskonałość działań", "Elita miasta", "Bezpieczeństwo"
        ],
        "questions": [
            "Jak komunikujesz mieszkańcom swoją pozycję lidera w Lublinie?",
            "W jaki sposób Twoje działania budują autorytet i zaufanie mieszkańców?",
            "Co robisz, by decyzje były stanowcze i jednoznaczne?",
            "Jak Twoje dokonania i inwestycje wzmacniają prestiż oraz bezpieczeństwo miasta?",
            "Jak zachęcasz wyborców do świadomego, silnego przywództwa?"
        ]
    },
    "Bohater": {
        "name": "Bohater",
        "tagline": "Determinacja. Odwaga. Sukces.",
        "description": (
            "Bohater w polityce to archetyp waleczności, determinacji i odwagi w podejmowaniu trudnych decyzji dla społeczności. "
            "Prezydent-Bohater mobilizuje mieszkańców do działania, bierze odpowiedzialność w najtrudniejszych momentach i broni interesów Lublina nawet pod presją."
        ),
        "storyline": (
            "Opowieść Bohatera to historia przezwyciężania kryzysów i stawania po stronie obywateli. "
            "Bohater nie rezygnuje nigdy, nawet w obliczu przeciwności. Jego postawa inspiruje i daje przykład innym samorządowcom."
        ),
        "recommendations": [
            "Komunikuj gotowość do działania, podkreślaj determinację w rozwiązywaniu problemów.",
            "Pokaż sukcesy i przykłady walki o interes mieszkańców.",
            "Stosuj dynamiczny język: zaznaczaj odwagę, mobilizację, sukces.",
            "Kolorystyka: czerwień, granat, biel.",
            "Pokazuj się w trudnych sytuacjach – reaguj natychmiast.",
            "Inspiruj współpracowników i mieszkańców do aktywności.",
            "Unikaj bierności, podkreślaj proaktywność."
        ],
        "core_traits": [
            "Odwaga", "Siła", "Determinacja", "Poświęcenie", "Sukces", "Inspiracja"
        ],
        "example_brands": [
            "Nike", "Polska Husaria", "ONG", "Patriotyczny samorząd"
        ],
        "color_palette": [
            "#E10600", "#2E3141", "#FFFFFF"
        ],
        "visual_elements": [
            "Peleryna", "Tarcza", "Aura odwagi", "Podniesiona dłoń", "Gwiazda"
        ],
        "keyword_messaging": [
            "Siła", "Zwycięstwo", "Poświęcenie", "Mobilizacja"
        ],
        "questions": [
            "Jak komunikujesz skuteczność w przezwyciężaniu kryzysów?",
            "Jak budujesz wizerunek walczącego o dobro mieszkańców?",
            "Jak pokazać determinację i niezłomność w działaniu publicznym?",
            "Które sukcesy świadczą o Twoim zaangażowaniu w trudnych sprawach?"
        ]
    },
    "Mędrzec": {
        "name": "Mędrzec",
        "tagline": "Wiedza. Racjonalność. Strategia.",
        "description": (
            "Mędrzec w polityce opiera komunikację na wiedzy, argumentacji i logicznym rozumowaniu. "
            "Kandydat na prezydenta wykorzystuje rozsądne analizy, doświadczenie oraz ekspercką wiedzę, by podejmować najlepsze decyzje dla całej społeczności."
        ),
        "storyline": (
            "Opowieść Mędrca to budowanie zaufania kompetencjami, przejrzystym uzasadnieniem propozycji i edukacją mieszkańców. "
            "Mędrzec nie działa pod wpływem impulsu; każda decyzja jest przemyślana i poparta faktami oraz wsłuchaniem się w potrzeby miasta."
        ),
        "recommendations": [
            "Wskazuj kompetencje, doświadczenie i eksperckość w zarządzaniu Lublinem.",
            "Komunikuj zrozumiale zawiłości miejskich inwestycji i decyzji.",
            "Stosuj wykresy, dane, analizy i argumenty – przemawiaj do rozumu obywateli.",
            "Zachowaj spokojny, opanowany ton.",
            "Używaj kolorystyki: błękit, szarość, granat.",
            "Podkreślaj racjonalność decyzji i transparentność działań.",
            "Unikaj populizmu – opieraj komunikację na faktach."
        ],
        "core_traits": [
            "Wiedza", "Rozwój", "Analiza", "Strategia", "Refleksja"
        ],
        "example_brands": [
            "Google", "Wikipedia", "MIT", "Think tanki"
        ],
        "color_palette": [
            "#4682B4", "#B0C4DE", "#6C7A89"
        ],
        "visual_elements": [
            "Okulary", "Księga", "Wykres", "Lupa", "Symbole nauki"
        ],
        "keyword_messaging": [
            "Wiedza", "Argument", "Racjonalność", "Rozwój miasta"
        ],
        "questions": [
            "Jak podkreślasz swoje doświadczenie i kompetencje?",
            "Jak przekonujesz mieszkańców argumentami i faktami?",
            "Jak edukujesz oraz tłumaczysz skomplikowane zmiany w mieście?",
            "W czym wyrażasz przewagę eksperckiej wiedzy nad populizmem?"
        ]
    },
    "Opiekun": {
        "name": "Opiekun",
        "tagline": "Empatia. Troska. Bezpieczeństwo.",
        "description": (
            "Opiekun w polityce to archetyp zaangażowania, wspierania i budowania poczucia wspólnoty. "
            "Kandydat–Opiekun dba o najsłabszych, promuje działania prospołeczne, wdraża programy pomocowe i społecznie odpowiedzialne."
        ),
        "storyline": (
            "Narracja Opiekuna podkreśla działania integrujące, troskę o seniorów, rodziny, niepełnosprawnych i osoby wykluczone. "
            "Buduje poczucie bezpieczeństwa oraz odpowiedzialności urzędu miasta za wszystkich obywateli."
        ),
        "recommendations": [
            "Akcentuj działania na rzecz integracji i wsparcia mieszkańców.",
            "Pokaż realne efekty programów prospołecznych i pomocowych.",
            "Stosuj ciepłą kolorystykę: zieleń, błękit, żółcie.",
            "Używaj symboliki: dłonie, serca, uścisk.",
            "Komunikuj empatię i autentyczną troskę o każdą grupę mieszkańców.",
            "Prowadź otwarte konsultacje społeczne.",
            "Unikaj twardego, technokratycznego tonu."
        ],
        "core_traits": [
            "Empatia", "Troska", "Wspólnota", "Bezpieczeństwo", "Solidarność"
        ],
        "example_brands": [
            "UNICEF", "Caritas", "WOŚP", "Hospicja"
        ],
        "color_palette": [
            "#B4D6B4", "#A7C7E7", "#FFD580"
        ],
        "visual_elements": [
            "Dłonie", "Serce", "Koło wspólnoty", "Symbol opieki"
        ],
        "keyword_messaging": [
            "Bezpieczeństwo mieszkańców", "Troska", "Wspólnota"
        ],
        "questions": [
            "Jak pokazujesz troskę i empatię wobec wszystkich mieszkańców?",
            "Jakie realne efekty mają wdrożone przez Ciebie programy pomocowe?",
            "W czym przejawia się Twoja polityka integrująca?",
            "Jak oceniasz skuteczność działań społecznych w mieście?"
        ]
    },
    "Kochanek": {
        "name": "Kochanek",
        "tagline": "Bliskość. Relacje. Pasja.",
        "description": (
            "Kochanek w polityce buduje pozytywne relacje z mieszkańcami, jest otwarty, komunikatywny i wzbudza zaufanie. "
            "Potrafi zbliżyć do siebie wyborców i sprawić, by czuli się zauważeni oraz docenieni."
        ),
        "storyline": (
            "Narracja Kochanka promuje serdeczność, ciepło i partnerskie traktowanie obywateli. "
            "Akcentuje jakość relacji z mieszkańcami, zespołem i innymi samorządami."
        ),
        "recommendations": [
            "Buduj relacje oparte na dialogu i wzajemnym szacunku.",
            "Stosuj ciepły, otwarty ton komunikacji.",
            "Promuj wydarzenia i inicjatywy integrujące społeczność.",
            "Używaj kolorystyki: czerwienie, róże, delikatne fiolety.",
            "Pokazuj, że wyborca jest dla Ciebie ważny.",
            "Doceniaj pozytywne postawy, sukcesy mieszkańców.",
            "Unikaj oficjalnego, zimnego tonu."
        ],
        "core_traits": [
            "Ciepło", "Relacje", "Bliskość", "Pasja", "Akceptacja"
        ],
        "example_brands": [
            "Allegro", "Santander", "Lubelskie Dni Róż"
        ],
        "color_palette": [
            "#FA709A", "#FEE140", "#FFD6E0"
        ],
        "visual_elements": [
            "Serce", "Uśmiech", "Gest bliskości"
        ],
        "keyword_messaging": [
            "Relacje", "Bliskość", "Społeczność"
        ],
        "questions": [
            "Jak komunikujesz otwartość i serdeczność wyborcom?",
            "Jakie działania podejmujesz, aby budować pozytywne relacje w mieście?",
            "Co robisz, by mieszkańcy czuli się ważni i zauważeni?"
        ]
    },
    "Błazen": {
        "name": "Błazen",
        "tagline": "Poczucie humoru. Dystans. Entuzjazm.",
        "description": (
            "Błazen w polityce wnosi lekkość, dystans i rozładowanie napięć. "
            "Kandydat-Błazen potrafi rozbawić, rozproszyć atmosferę, ale nigdy nie traci dystansu do siebie i powagi spraw publicznych."
        ),
        "storyline": (
            "Narracja Błazna to umiejętność śmiania się z problemów i codziennych wyzwań miasta, ale też dawania mieszkańcom nadziei oraz pozytywnej energii."
        ),
        "recommendations": [
            "Stosuj humor w komunikacji (ale z umiarem i klasą!).",
            "Rozluźniaj atmosferę podczas spotkań i debat.",
            "Podkreślaj pozytywne aspekty życia w mieście.",
            "Kolorystyka: żółcie, pomarańcze, intensywne kolory.",
            "Nie bój się autoironii.",
            "Promuj wydarzenia integrujące, rozrywkowe.",
            "Unikaj przesadnego formalizmu."
        ],
        "core_traits": [
            "Poczucie humoru", "Entuzjazm", "Dystans", "Optymizm"
        ],
        "example_brands": [
            "Allegro", "Łomża", "Kabarety"
        ],
        "color_palette": [
            "#FFB300", "#FF8300", "#FFD93D"
        ],
        "visual_elements": [
            "Uśmiech", "Czapka błazna", "Kolorowe akcenty"
        ],
        "keyword_messaging": [
            "Dystans", "Entuzjazm", "Radość"
        ],
        "questions": [
            "W jaki sposób wykorzystujesz humor w komunikacji publicznej?",
            "Jak rozładowujesz napięcia w sytuacjach kryzysowych?",
            "Co robisz, aby mieszkańcy mogli wspólnie się bawić i śmiać?"
        ]
    },
    "Twórca": {
        "name": "Twórca",
        "tagline": "Kreatywność. Innowacja. Wizja.",
        "description": (
            "Twórca w polityce to źródło nowych pomysłów, innowacji i niebanalnych rozwiązań dla miasta. "
            "Jako prezydent–Twórca nie boi się wdrażać oryginalnych, często nieszablonowych strategii."
        ),
        "storyline": (
            "Opowieść Twórcy jest oparta na zmianie, wprowadzaniu kreatywnych rozwiązań oraz inspirowaniu innych do współdziałania dla rozwoju Lublina."
        ),
        "recommendations": [
            "Proponuj i wdrażaj nietypowe rozwiązania w mieście.",
            "Pokazuj przykłady innowacyjnych projektów.",
            "Promuj kreatywność i otwartość na zmiany.",
            "Stosuj kolorystykę: zielenie, lazurowe błękity, fiolety.",
            "Doceniaj artystów, startupy, lokalne inicjatywy.",
            "Buduj wizerunek miasta-innowatora.",
            "Unikaj schematów i powtarzalnych projektów."
        ],
        "core_traits": [
            "Kreatywność", "Odwaga twórcza", "Inspiracja", "Wizja", "Nowatorstwo"
        ],
        "example_brands": [
            "Tesla", "Dyrekcja Lublin", "Startupy"
        ],
        "color_palette": [
            "#7C53C3", "#3BE8B0", "#87CEEB"
        ],
        "visual_elements": [
            "Kostka Rubika", "Żarówka", "Kolorowe fale"
        ],
        "keyword_messaging": [
            "Innowacja", "Twórczość", "Wizja rozwoju"
        ],
        "questions": [
            "Jak promujesz kreatywność i innowacyjność w mieście?",
            "Jakie oryginalne projekty wdrożyłeś lub planujesz wdrożyć?",
            "Jak inspirować mieszkańców do kreatywnego działania?"
        ]
    },
    "Odkrywca": {
        "name": "Odkrywca",
        "tagline": "Odwaga. Ciekawość. Nowe horyzonty.",
        "description": (
            "Odkrywca poszukuje nowych rozwiązań, jest otwarty na zmiany i śledzi światowe trendy, które wdraża w Lublinie. "
            "Wybiera nowatorskie, nieoczywiste drogi dla rozwoju miasta i jego mieszkańców."
        ),
        "storyline": (
            "Opowieść Odkrywcy to wędrowanie poza schematami, miasto bez barier, eksperymentowanie z nowościami oraz angażowanie mieszkańców w odkrywcze projekty."
        ),
        "recommendations": [
            "Inicjuj nowe projekty i szukaj innowacji także poza Polską.",
            "Promuj przełamywanie standardów i aktywność obywatelską.",
            "Stosuj kolorystykę: turkusy, błękity, odcienie zieleni.",
            "Publikuj inspiracje z innych miast i krajów.",
            "Wspieraj wymiany młodzieży, startupy, koła naukowe.",
            "Unikaj stagnacji i powielania dawnych schematów."
        ],
        "core_traits": [
            "Odwaga", "Ciekawość", "Niezależność", "Nowatorstwo"
        ],
        "example_brands": [
            "Red Bull", "National Geographic", "Nomadzi"
        ],
        "color_palette": [
            "#43C6DB", "#A0E8AF", "#F9D371"
        ],
        "visual_elements": [
            "Mapa", "Kompas", "Droga", "Lupa"
        ],
        "keyword_messaging": [
            "Odkrywanie", "Nowe horyzonty", "Zmiana"
        ],
        "questions": [
            "Jak zachęcasz do odkrywania nowości w mieście?",
            "Jakie projekty wdrażasz, które nie były jeszcze realizowane w innych miastach?",
            "Jak budujesz wizerunek Lublina jako miejsca wolnego od barier?"
        ]
    },
    "Czarodziej": {
        "name": "Czarodziej",
        "tagline": "Transformacja. Inspiracja. Przełom.",
        "description": (
            "Czarodziej w polityce to wizjoner i transformator – wytycza nowy kierunek i inspiruje do zmian niemożliwych na pierwszy rzut oka. "
            "Dzięki jego inicjatywom Lublin przechodzi metamorfozy, w których niemożliwe staje się możliwe."
        ),
        "storyline": (
            "Opowieść Czarodzieja to zmiana wykraczająca poza rutynę, wyobraźnia, inspiracja, a także odwaga w stawianiu pytań i szukaniu odpowiedzi poza schematami."
        ),
        "recommendations": [
            "Wprowadzaj śmiałe, czasem kontrowersyjne pomysły w życie.",
            "Podkreślaj rolę wizji i inspiracji.",
            "Stosuj symbolikę: gwiazdy, zmiany, światło, 'magiczne' efekty.",
            "Stosuj kolorystykę: fiolety, granaty, akcent perłowy.",
            "Buduj wyobrażenie miasta jako miejsca możliwości.",
            "Unikaj banalnych, powtarzalnych rozwiązań."
        ],
        "core_traits": [
            "Inspiracja", "Przemiana", "Wyobraźnia", "Transcendencja"
        ],
        "example_brands": [
            "Apple", "Disney", "Nasa", "Nowoczesny Lublin"
        ],
        "color_palette": [
            "#8F00FF", "#181C3A", "#E0BBE4"
        ],
        "visual_elements": [
            "Gwiazda", "Iskra", "Łuk magiczny"
        ],
        "keyword_messaging": [
            "Zmiana", "Inspiracja", "Możliwość"
        ],
        "questions": [
            "Jak pokazujesz mieszkańcom, że niemożliwe jest możliwe?",
            "Jakie innowacje budują wizerunek miasta kreatywnego i nowoczesnego?",
            "Jak inspirujesz społeczność do patrzenia dalej?"
        ]
    },
    "Towarzysz": {
        "name": "Towarzysz",
        "tagline": "Wspólnota. Prostota. Bliskość.",
        "description": (
            "Towarzysz w polityce stoi blisko ludzi, jest autentyczny, stawia na prostotę i tworzenie bezpiecznej wspólnoty społecznej. "
            "Nie udaje, nie buduje dystansu – jest 'swojakiem', na którym można polegać."
        ),
        "storyline": (
            "Opowieść Towarzysza koncentruje się wokół wartości rodzinnych, codziennych wyzwań, pracy od podstaw oraz pielęgnowania lokalnej tradycji."
        ),
        "recommendations": [
            "Podkreślaj prostotę i codzienność w komunikacji.",
            "Stosuj jasne, proste słowa i obrazy.",
            "Buduj atmosferę równości (każdy ma głos).",
            "Stosuj kolorystykę: beże, błękity, zielone akcenty.",
            "Doceniaj lokalność i rodzinność.",
            "Promuj wspólnotowe inicjatywy.",
            "Unikaj dystansu i języka eksperckiego."
        ],
        "core_traits": [
            "Autentyczność", "Wspólnota", "Prostota", "Równość"
        ],
        "example_brands": [
            "Sieci osiedlowe", "Społem", "Allegro"
        ],
        "color_palette": [
            "#F9F9F9", "#6CA0DC", "#A3C1AD"
        ],
        "visual_elements": [
            "Dom", "Krąg ludzi", "Prosta ikona dłoni"
        ],
        "keyword_messaging": [
            "Bliskość", "Razem", "Prostota"
        ],
        "questions": [
            "Jak podkreślasz autentyczność i codzienność?",
            "Jak pielęgnujesz lokalność i wspólnotę?",
            "Co robisz, by każdy mieszkaniec czuł się zauważony?"
        ]
    },
    "Niewinny": {
        "name": "Niewinny",
        "tagline": "Optymizm. Nadzieja. Nowy początek.",
        "description": (
            "Niewinny w polityce otwarcie komunikuje pozytywne wartości, niesie nadzieję i podkreśla wiarę w zmiany na lepsze. "
            "Kandydat–Niewinny buduje zaufanie szczerością i skutecznie apeluje o współpracę dla wspólnego dobra."
        ),
        "storyline": (
            "Opowieść Niewinnego buduje napięcie wokół pozytywnych emocji, odwołuje się do marzeń o lepszym Lublinie i wiary we wspólny sukces."
        ),
        "recommendations": [
            "Komunikuj optymizm, wiarę w ludzi i dobre intencje.",
            "Stosuj jasną kolorystykę: biele, pastele, żółcie.",
            "Dziel się sukcesami społeczności.",
            "Stawiaj na transparentność działań.",
            "Angażuj się w kampanie edukacyjne i społeczne.",
            "Unikaj negatywnego przekazu, straszenia, manipulacji."
        ],
        "core_traits": [
            "Optymizm", "Nadzieja", "Współpraca", "Szlachetność"
        ],
        "example_brands": [
            "Kinder", "Polska Akcja Humanitarna"
        ],
        "color_palette": [
            "#FFF6C3", "#AAC9CE", "#FFF200"
        ],
        "visual_elements": [
            "Gołąb", "Słońce", "Dziecko"
        ],
        "keyword_messaging": [
            "Nadzieja", "Optymizm", "Wspólnie"
        ],
        "questions": [
            "Jak budujesz wizerunek pozytywnego samorządowca?",
            "Jak zachęcasz mieszkańców do dzielenia się nadzieją?",
            "Jak komunikujesz szczerość i otwartość?"
        ]
    },
    "Buntownik": {
        "name": "Buntownik",
        "tagline": "Zmiana. Odwaga. Przełom.",
        "description": (
            "Buntownik w polityce odważnie kwestionuje zastane układy, nawołuje do zmiany i walczy o nowe, lepsze reguły gry w mieście. "
            "Potrafi ściągnąć uwagę i zjednoczyć mieszkańców wokół śmiałych idei."
        ),
        "storyline": (
            "Narracja Buntownika podkreśla walkę z niesprawiedliwością i stagnacją, wytykanie błędów władzy i radykalne pomysły na rozwój Lublina."
        ),
        "recommendations": [
            "Akcentuj odwagę do mówienia „nie” starym rozwiązaniom.",
            "Publikuj manifesty i odważne postulaty.",
            "Stosuj wyrazistą kolorystykę: czernie, czerwienie, ostre kolory.",
            "Inspiruj mieszkańców do aktywnego sprzeciwu wobec barier rozwojowych.",
            "Podkreślaj wolność słowa, swobody obywatelskie.",
            "Unikaj koncentrowania się wyłącznie na krytyce – pokazuj pozytywne rozwiązania."
        ],
        "core_traits": [
            "Odwaga", "Bezpardonowość", "Radykalizm", "Niepokorność"
        ],
        "example_brands": [
            "Harley Davidson", "Greenpeace", "Gazeta Wyborcza"
        ],
        "color_palette": [
            "#000000", "#FF0000", "#FF6F61"
        ],
        "visual_elements": [
            "Piorun", "Megafon", "Odwrócona korona"
        ],
        "keyword_messaging": [
            "Zmiana", "Rewolucja", "Nowe reguły"
        ],
        "questions": [
            "Jak komunikujesz odwagę i gotowość do zmiany?",
            "Jak mobilizujesz do zrywania z przeszłością?",
            "Co robisz, by mieszkańcy mieli w Tobie rzecznika zmiany?"
        ]
    }
}

ARCHE_IMG_URL = "https://justynakopec.pl/wp-content/uploads/2024/08/Archetypy-marki-Justyna-Kopec.png"
ARCHE_NAME_TO_IDX = {n.lower(): i for i, n in enumerate(ARCHE_NAMES_ORDER)}

@st.cache_data
def load_base_arche_img():
    resp = requests.get(ARCHE_IMG_URL, stream=True, timeout=30)
    resp.raise_for_status()
    img = Image.open(io.BytesIO(resp.content)).convert("RGBA")
    return img

def mask_for(idx, color):
    base = load_base_arche_img()
    w, h = base.size
    cx, cy = w//2, h//2
    rad = w//2
    mask = Image.new("RGBA", (w, h), (0,0,0,0))
    draw = ImageDraw.Draw(mask)
    start = -90 + idx*30
    end = start + 30
    draw.pieslice([cx-rad, cy-rad, cx+rad, cy+rad], start, end, fill=color)
    return mask

def compose_archetype_highlight(idx_main, idx_aux=None):
    base = load_base_arche_img().copy()
    if idx_aux is not None and idx_aux != idx_main and idx_aux < 12:
        mask_aux = mask_for(idx_aux, (255,210,47,140))
        base.alpha_composite(mask_aux)
    if idx_main is not None:
        mask_main = mask_for(idx_main, (255,0,0,140))
        base.alpha_composite(mask_main)
    return base

def archetype_name_to_img_idx(name):
    try:
        return ARCHE_NAMES_ORDER.index(name)
    except ValueError:
        return None

def clean_pdf_text(text):
    if text is None:
        return ""
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = text.replace("–", "-").replace("—", "-")
    return text

@st.cache_data(ttl=30)
def load():
    try:
        conn = psycopg2.connect(
            host=st.secrets["db_host"],
            database=st.secrets["db_name"],
            user=st.secrets["db_user"],
            password=st.secrets["db_pass"],
            port=st.secrets.get("db_port", 5432),
            sslmode="require"
        )
        df = pd.read_sql("SELECT * FROM ap48_responses", con=conn)
        conn.close()
        if "created_at" in df.columns:
            df["created_at"] = pd.to_datetime(df["created_at"])
        def parse_answers(x):
            if isinstance(x, list):
                return x
            try:
                import json
                return json.loads(x)
            except:
                try:
                    return ast.literal_eval(x)
                except:
                    return None
        if "answers" in df.columns:
            df["answers"] = df["answers"].apply(parse_answers)
        return df
    except Exception as e:
        st.warning(f"Błąd podczas ładowania danych: {e}")
        return pd.DataFrame()

def archetype_scores(answers):
    if not isinstance(answers, list) or len(answers) < 48:
        return {k: None for k in archetypes}
    out = {}
    for name, idxs in archetypes.items():
        out[name] = sum(answers[i-1] for i in idxs)
    return out

def archetype_percent(scoresum):
    if scoresum is None:
        return None
    return round(scoresum / 20 * 100, 1)

def pick_main_and_aux_archetype(archetype_means, archetype_order):
    vals = list(archetype_means.values())
    max_val = max(vals)
    main_candidates = [k for k, v in archetype_means.items() if v == max_val]
    main_type = next(k for k in archetype_order if k in main_candidates)
    aux_vals = [v for k, v in archetype_means.items() if k != main_type]
    if not aux_vals:
        return main_type, None
    aux_val = max(aux_vals)
    aux_candidates = [k for k, v in archetype_means.items() if v == aux_val and k != main_type]
    second_type = next((k for k in archetype_order if k in aux_candidates), None)
    return main_type, second_type

def export_word(main_type, second_type, features, main, second):
    doc = Document()
    doc.add_heading("Raport AP-48 – Archetypy", 0)
    doc.add_heading(f"Główny archetyp: {main_type}", level=1)
    doc.add_paragraph(f"Cechy kluczowe: {features.get(main_type, '-')}")
    doc.add_paragraph(main.get("description", ""))
    doc.add_paragraph("Storyline: " + main.get("storyline", ""))
    doc.add_paragraph("Rekomendacje: " + "\n".join(main.get("recommendations", [])))
    if second_type and second_type != main_type:
        doc.add_heading(f"Archetyp pomocniczy: {second_type}", level=2)
        doc.add_paragraph(f"Cechy kluczowe: {features.get(second_type, '-')}")
        doc.add_paragraph(second.get("description", ""))
        doc.add_paragraph("Storyline: " + second.get("storyline", ""))
        doc.add_paragraph("Rekomendacje: " + "\n".join(second.get("recommendations", [])))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_pdf(main_type, second_type, features, main, second):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, clean_pdf_text("Raport AP-48 – Archetypy"), ln=1)
    pdf.set_font("Arial", "B", 13)
    pdf.cell(0, 8, clean_pdf_text(f"Główny archetyp: {main_type}"), ln=1)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(
        0, 7, clean_pdf_text(
            f"Cechy kluczowe: {features.get(main_type, '-')}\n\n"
            f"{main.get('description', '')}\n\n"
            f"Storyline: {main.get('storyline', '')}\n\n"
            f"Rekomendacje: " + "\n".join(main.get("recommendations", [])) + "\n"
        )
    )
    if second_type and second_type != main_type:
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, clean_pdf_text(f"Archetyp pomocniczy: {second_type}"), ln=1)
        pdf.set_font("Arial", "", 11)
        pdf.multi_cell(
            0, 7, clean_pdf_text(
                f"Cechy kluczowe: {features.get(second_type, '-')}\n\n"
                f"{second.get('description', '')}\n\n"
                f"Storyline: {second.get('storyline', '')}\n\n"
                f"Rekomendacje: " + "\n".join(second.get("recommendations", []))
            )
        )
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    buf = BytesIO(pdf_bytes)
    buf.seek(0)
    return buf

def is_color_dark(color_hex):
    if color_hex is None:
        return False
    if not color_hex.startswith('#') or len(color_hex) not in (7, 4):
        return False
    h = color_hex.lstrip('#')
    if len(h) == 3:
        h = ''.join([c*2 for c in h])
    r, g, b = tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
    lum = 0.2126*r + 0.7152*g + 0.0722*b
    return lum < 110

def render_archetype_card(archetype_data, main=True):
    if not archetype_data:
        st.warning("Brak danych o archetypie.")
        return
    border_color = archetype_data.get('color_palette', ['#888'])[0]
    tagline = archetype_data.get('tagline','')
    if (archetype_data.get('name','').strip().lower() == 'niewinny') and not main:
        tagline = "Niesie nadzieję, inspiruje do współpracy, buduje zaufanie szczerością i apeluje o wspólne dobro, otwarcie komunikuje pozytywne wartości."
    symbol = archetype_data.get('visual_elements', [''])[0] if archetype_data.get('visual_elements') else ""
    symbol_emoji = {
        "Korona": "👑", "Herb Lublina": "🛡️", "Peleryna": "🦸", "Serce": "❤️","Uśmiech": "😊","Dłonie": "🤝",
        "Księga": "📖", "Mapa": "🗺️","Gwiazda": "⭐", "Gołąb": "🕊️","Piorun": "⚡", "Rubika": "🧩", "Dom": "🏡"
    }
    icon = symbol_emoji.get(symbol, "🔹")
    box_shadow = f"0 4px 14px 0 {border_color}44" if main else f"0 2px 6px 0 {border_color}22"
    bg_color = "#FAFAFA" if not main else (archetype_data.get('color_palette', ['#FFF', '#FAFAFA'])[1])
    width_card = "70vw"
    text_color = "#222"
    tagline_color = "#88894A" if archetype_data.get('name','').lower() == "niewinny" else border_color
    if main and is_color_dark(bg_color):
        text_color = "#fff"
        tagline_color = "#FFD22F" if archetype_data.get('name','').lower() == "bohater" else "#fffbea"
    st.markdown(f"""
    <div style="
        max-width:{width_card};
        border: 3px solid {border_color if main else '#CCC'};
        border-radius: 20px;
        background: {bg_color};
        box-shadow: {box_shadow};
        padding: 2.1em 2.2em 1.3em 2.2em;
        margin-bottom: 16px;
        color: {text_color};
        display: flex; align-items: flex-start;">
        <div style="font-size:2.6em; margin-right:23px; margin-top:3px; flex-shrink:0;">
            {icon}
        </div>
        <div>
            <div style="font-size:2.15em;font-weight:bold; line-height:1.08; margin-bottom:1px; color:{text_color};">
                {archetype_data.get('name','?')}
            </div>
            <div style="font-size:1.22em; font-style:italic; color:{tagline_color}; margin-bottom:18px; margin-top:4px;">
                {tagline}
            </div>
            <div style="margin-top:21px;"><b>Opis:</b> {archetype_data.get('description','')}</div>
            <div style="margin-top:24px;font-weight:600;">Storyline:</div>
            <div style="margin-bottom:9px; margin-top:4px;">{archetype_data.get('storyline','')}</div>
            <div style="color:#666;font-size:1em; margin-top:21px;"><b>Cechy:</b> {", ".join(archetype_data.get('core_traits',[]))}</div>
            <div style="margin-top:24px;font-weight:600;">Rekomendacje:</div>
            <ul style="padding-left:24px; margin-bottom:9px;">
                {''.join(f'<li style="margin-bottom:2px;">{r}</li>' for r in archetype_data.get('recommendations',[]))}
            </ul>
            <div style="margin-top:29px;font-weight:600;">Słowa kluczowe:</div>
            <div style="margin-bottom:8px;">{', '.join(archetype_data.get('keyword_messaging',[]))}</div>
            <div style="margin-top:24px;font-weight:600;">Elementy wizualne:</div>
            <div style="margin-bottom:8px;">{', '.join(archetype_data.get('visual_elements',[]))}</div>
            <div style="margin-top:24px;font-weight:600;">Przykłady marek/organizacji:</div>
            <div>{', '.join(archetype_data.get('example_brands',[]))}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

data = load()
num_ankiet = len(data) if not data.empty else 0

header_col1, header_col2 = st.columns([0.77, 0.23])
with header_col1:
    st.markdown("""
    <div style="font-size:2.3em; font-weight:bold; background:#1a93e3; color:#fff; 
        padding:14px 32px 10px 24px; border-radius:2px; width:fit-content; display:inline-block;">
        Archetypy Krzysztofa Hetmana – panel administratora
    </div>
    """, unsafe_allow_html=True)
with header_col2:
    st.markdown(f"""
    <div style="display:flex;align-items:center;justify-content:flex-end;height:100%;"><div style="font-size:1.23em;text-align:right;background:#f3f3fa;padding:12px 29px 8px 29px; border-radius:17px; border:2px solid #d1d9ed;color:#195299;font-weight:600;box-shadow:0 2px 10px 0 #b5c9e399;">
        <span style="font-size:1.8em;font-weight:bold;">{num_ankiet}</span><br/>uczestników badania
    </div></div>
    """, unsafe_allow_html=True)

st.markdown("""
<hr style="height:1.3px;background:#eaeaec; margin-top:1.8em; margin-bottom:3.8em; border:none;" />
""", unsafe_allow_html=True)

if "answers" in data.columns and not data.empty:
    results = []
    for idx, row in data.iterrows():
        if not isinstance(row.get("answers", None), list):
            continue
        arcsums = archetype_scores(row["answers"])
        arcper = {k: archetype_percent(v) for k, v in arcsums.items()}
        main_type, second_type = pick_main_and_aux_archetype(arcsums, ARCHE_NAMES_ORDER)
        main = archetype_extended.get(main_type, {})
        second = archetype_extended.get(second_type, {}) if second_type != main_type else {}
        results.append({
            "ID": row.get("id", idx+1),
            **arcsums,
            **{f"{k}_%" : v for k,v in arcper.items()},
            "Główny archetyp": main_type,
            "Cechy kluczowe": archetype_features.get(main_type,""),
            "Opis": main.get("description", ""),
            "Storyline": main.get("storyline", ""),
            "Rekomendacje": "\n".join(main.get("recommendations", [])),
            "Archetyp pomocniczy": second_type if second_type != main_type else "",
            "Cechy pomocniczy": archetype_features.get(second_type,"") if second_type != main_type else "",
            "Opis pomocniczy": second.get("description", "") if second_type != main_type else "",
            "Storyline pomocniczy": second.get("storyline", "") if second_type != main_type else "",
            "Rekomendacje pomocniczy": "\n".join(second.get("recommendations", [])) if second_type != main_type else "",
        })
    results_df = pd.DataFrame(results)

    if not results_df.empty and "ID" in results_df.columns:
        results_df = results_df.sort_values("ID")
        archetype_names = ARCHE_NAMES_ORDER
        counts_main = results_df['Główny archetyp'].value_counts().reindex(archetype_names, fill_value=0)
        counts_aux = results_df['Archetyp pomocniczy'].value_counts().reindex(archetype_names, fill_value=0)
        archetype_emoji = {
            "Władca":"👑", "Bohater":"🦸", "Mędrzec":"📖", "Opiekun":"🤝", "Kochanek":"❤️",
            "Błazen":"😂", "Twórca":"🧩", "Odkrywca":"🗺️", "Czarodziej":"⭐", "Towarzysz":"🏡",
            "Niewinny":"🕊️", "Buntownik":"⚡"
        }
        def zero_to_dash(val): return "-" if val == 0 else str(val)
        archetype_table = pd.DataFrame({
            "Archetyp": [f"{archetype_emoji.get(n,n)} {n}" for n in archetype_names],
            "Główny archetyp": [zero_to_dash(counts_main.get(k, 0)) for k in archetype_names],
            "Pomocniczy archetyp": [zero_to_dash(counts_aux.get(k, 0)) for k in archetype_names]
        })
        archetype_table_html = archetype_table.to_html(escape=False, index=False)
        archetype_table_html = archetype_table_html.replace('<th>', '<th style="text-align:center">')
        archetype_table_html = archetype_table_html.replace('<td>', '<td style="text-align:center">')


        def align_first_column_to_left_with_width(html):
            html = re.sub(
                r'(<tr[^>]*>)(\s*<td style="text-align:center">)',
                lambda m: m.group(1) + m.group(2).replace('text-align:center', 'text-align:left;width:24%;'), html
            )
            html = html.replace(
                '<th style="text-align:center">Archetyp</th>',
                '<th style="text-align:center;width:24%;">Archetyp</th>'
            )
            html = html.replace(
                '<th style="text-align:center">Główny archetyp</th>',
                '<th style="text-align:center;width:18%;">Główny archetyp</th>'
            ).replace(
                '<th style="text-align:center">Pomocniczy archetyp</th>',
                '<th style="text-align:center;width:18%;">Pomocniczy archetyp</th>'
            )
            html = re.sub(
                r'<tr>(\s*<td style="[^"]*left;?[^"]*">.*?</td>)'
                r'(\s*<td style="text-align:center">)',
                r'<tr>\1<td style="text-align:center;width:18%">', html
            )
            html = re.sub(
                r'(<td style="text-align:center;width:18%">.*?</td>)'
                r'(\s*<td style="text-align:center">)',
                r'\1<td style="text-align:center;width:18%">', html
            )
            return html
        archetype_table_html = align_first_column_to_left_with_width(archetype_table_html)
        archetype_table_html = archetype_table_html.replace(
            '<table ',
            '<table style="margin-left:0px;margin-right:0px;width:99%;" '
        )
        st.markdown('<div style="font-size:2.1em;font-weight:600;margin-bottom:22px;">Informacje na temat archetypu Krzysztofa Hetmana</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns([0.23, 0.40, 0.42], gap="small")
        with col1:
            st.markdown('<div style="font-size:1.3em;font-weight:600;margin-bottom:13px;">Liczebność archetypów głównych i pomocniczych</div>', unsafe_allow_html=True)
            st.markdown(archetype_table_html, unsafe_allow_html=True)
        with col2:
            mean_archetype_scores = {k: results_df[k].mean() if k in results_df.columns else 0 for k in archetype_names}
            main_type, second_type = pick_main_and_aux_archetype(mean_archetype_scores, archetype_names)
            theta_labels = []
            for n in archetype_names:
                if n == main_type:
                    theta_labels.append(f"<b><span style='color:red;'>{n}</span></b>")
                elif n == second_type:
                    theta_labels.append(f"<b><span style='color:#FFD22F;'>{n}</span></b>")
                else:
                    theta_labels.append(f"<span style='color:#656565;'>{n}</span>")
            highlight_r = []
            highlight_marker_color = []
            for i, name in enumerate(archetype_names):
                if name == main_type:
                    highlight_r.append(mean_archetype_scores[name])
                    highlight_marker_color.append("red")
                elif name == second_type:
                    highlight_r.append(mean_archetype_scores[name])
                    highlight_marker_color.append("#FFD22F")
                else:
                    highlight_r.append(None)
                    highlight_marker_color.append("rgba(0,0,0,0)")
            st.markdown('<div style="font-size:1.3em;font-weight:600;margin-bottom:13px; text-align:center;">Profil archetypów Krzysztofa Hetmana</div>', unsafe_allow_html=True)
            fig = go.Figure(
                data=[
                    go.Scatterpolar(
                        r=list(mean_archetype_scores.values()) + [list(mean_archetype_scores.values())[0]],
                        theta=archetype_names + [archetype_names[0]],
                        fill='toself',
                        name='Średnia wszystkich',
                        line=dict(color="royalblue", width=3),
                        marker=dict(size=6)
                    ),
                    go.Scatterpolar(
                        r=highlight_r,
                        theta=archetype_names,
                        mode='markers',
                        marker=dict(size=18, color=highlight_marker_color, opacity=0.95, line=dict(color="black", width=2)),
                        name='Archetyp główny / pomocniczy',
                        showlegend=False,
                    )
                ],
                layout=go.Layout(
                    polar=dict(
                        radialaxis=dict(visible=True, range=[0, 20]),
                        angularaxis=dict(tickfont=dict(size=19), tickvals=archetype_names, ticktext=theta_labels)
                    ),
                    width=550, height=550,
                    margin=dict(l=20, r=20, t=32, b=32),
                    showlegend=False
                )
            )
            st.plotly_chart(fig, use_container_width=True)
            st.markdown("""
            <div style="display: flex; justify-content: center; align-items: center; margin:10px 0 3px 0;">
            <span style="display:inline-block;vertical-align:middle;width:21px;height:21px;border-radius:50%;background:red;border:2px solid black;margin-right:6px"></span>
            <span style="font-size:0.80em;vertical-align:middle;margin-right:18px; color:#111;">Archetyp główny</span>
            <span style="display:inline-block;vertical-align:middle;width:21px;height:21px;border-radius:50%;background:#FFD22F;border:2px solid black;margin-right:6px"></span>
            <span style="font-size:0.80em;vertical-align:middle;color:#555;">Archetyp pomocniczy</span>
            </div>
            """, unsafe_allow_html=True)
        with col3:
            if main_type is not None:
                kola_img = compose_archetype_highlight(
                    archetype_name_to_img_idx(main_type),
                    archetype_name_to_img_idx(second_type) if second_type != main_type else None
                )
                st.image(
                    kola_img,
                    caption="Podświetlenie: główny – czerwony, pomocniczy – żółty",
                    width=700  # ← tu wpisujesz dokładnie, ile chcesz pikseli szerokości (np. 450, 512, 600)
                )

        st.markdown("""
        <hr style="height:1px; border:none; background:#eee; margin-top:34px; margin-bottom:19px;" />
        """, unsafe_allow_html=True)

        st.markdown(f'<div style="font-size:2.1em;font-weight:700;margin-bottom:16px;">Archetyp główny Krzysztofa Hetmana</div>', unsafe_allow_html=True)
        render_archetype_card(archetype_extended.get(main_type, {}), main=True)

        if second_type and second_type != main_type:
            st.markdown("<div style='height:12px;'></div>", unsafe_allow_html=True)
            st.markdown("""
            <hr style="height:1.1px; border:none; background:#ddd; margin-top:6px; margin-bottom:18px;" />
            """, unsafe_allow_html=True)
            st.markdown("<div style='font-size:1.63em;font-weight:700;margin-bottom:15px;'>Archetyp pomocniczy Krzysztofa Hetmana</div>", unsafe_allow_html=True)
            render_archetype_card(archetype_extended.get(second_type, {}), main=False)

        st.markdown("""
        <hr style="height:1px; border:none; background:#eee; margin-top:38px; margin-bottom:24px;" />
        """, unsafe_allow_html=True)

        st.markdown('<div style="font-size:1.13em;font-weight:600;margin-bottom:13px;">Tabela odpowiedzi respondentów (pełne wyniki)</div>', unsafe_allow_html=True)

        # --- FINALNY BLOK: TABELA Z PODSUMOWANIEM I EKSPORTEM XLSX ---

        final_df = results_df.copy()

        try:
            col_to_exclude = [
                "ID", "Archetyp", "Główny archetyp", "Cechy kluczowe", "Opis", "Storyline",
                "Rekomendacje", "Archetyp pomocniczy", "Cechy pomocniczy", "Opis pomocniczy",
                "Storyline pomocniczy", "Rekomendacje pomocniczy"
            ]
            means = final_df.drop(columns=col_to_exclude, errors="ignore").mean(numeric_only=True)
            summary_row = {col: round(means[col], 2) if col in means else "-" for col in final_df.columns}
            summary_row["ID"] = "ŚREDNIA"
            final_df = final_df.iloc[::-1]  # Odwróć kolejność wierszy (najnowsze na górze)
            final_df = pd.concat([final_df, pd.DataFrame([summary_row])], ignore_index=True)  # Dodaj średnią na koniec
        except Exception as e:
            pass  # Jeśli coś pójdzie nie tak, po prostu nie dodawaj podsumowania

        st.dataframe(final_df, hide_index=True)

        # Eksport CSV
        st.download_button("Pobierz wyniki archetypów (CSV)", final_df.to_csv(index=False), "ap48_archetypy.csv")

        # Eksport XLSX
        import io

        buffer = io.BytesIO()
        final_df.to_excel(buffer, index=False)
        st.download_button(
            label="Pobierz wyniki archetypów (XLSX)",
            data=buffer.getvalue(),
            file_name="ap48_archetypy.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
else:
    st.info("Brak danych 'answers' – nie wykryto odpowiedzi w bazie danych.")